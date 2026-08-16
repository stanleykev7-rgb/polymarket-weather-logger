"""
Performance report generation for the Weather Market Data Explorer.

Builds a "day-by-day model performance across lead times" report as
either a PNG image (quick visual, one page worth of charts) or a DOCX
document (more detail, includes data tables and narrative insights),
for either all cities combined or a single city.

This module is intentionally decoupled from Streamlit -- it takes a
plain pandas DataFrame in and returns raw bytes out -- so it can be
tested and reasoned about without spinning up the app itself.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone

import matplotlib
matplotlib.use("Agg")  # no display backend needed/available server-side
import matplotlib.pyplot as plt
import pandas as pd

LEAD_TIME_BINS = [-float("inf"), 24, 48, 72, 168, float("inf")]
LEAD_TIME_LABELS = ["<24h", "24-48h", "48-72h", "3-7d", "7d+"]

MODEL_HIT_COLS = {
    "ECMWF": "ecmwf_hit",
    "GFS": "gfs_hit",
    "Market Favorite": "market_favorite_hit",
}


# ---------------------------------------------------------------------------
# Data prep
# ---------------------------------------------------------------------------
def _settled_subset(df: pd.DataFrame, city: str | None) -> pd.DataFrame:
    """Rows with at least one resolved hit flag, optionally filtered to
    one city. Doesn't require ALL three hit columns to be non-null --
    older rows may only have ecmwf_hit populated (gfs_hit/
    market_favorite_hit didn't exist pre-schema-v2), and that's fine;
    each metric below handles its own column's nulls independently.
    """
    hit_cols = [c for c in MODEL_HIT_COLS.values() if c in df.columns]
    if not hit_cols:
        return df.iloc[0:0]
    mask = pd.Series(False, index=df.index)
    for c in hit_cols:
        mask = mask | df[c].isin([True, False])
    out = df[mask].copy()
    if city:
        out = out[out["city"] == city]
    return out


def compute_daily_performance(df: pd.DataFrame, city: str | None = None) -> pd.DataFrame:
    """One row per target_date: hit rate for each model, plus sample size."""
    settled = _settled_subset(df, city)
    if settled.empty:
        return pd.DataFrame(columns=["target_date", "n"] + [f"{m}_hit_rate" for m in MODEL_HIT_COLS])

    rows = []
    for target_date, group in settled.groupby("target_date"):
        row = {"target_date": target_date, "n": len(group)}
        for model, col in MODEL_HIT_COLS.items():
            if col in group.columns:
                vals = group[col].dropna()
                row[f"{model}_hit_rate"] = vals.mean() if len(vals) else None
                row[f"{model}_n"] = len(vals)
            else:
                row[f"{model}_hit_rate"] = None
                row[f"{model}_n"] = 0
        rows.append(row)
    out = pd.DataFrame(rows).sort_values("target_date")
    return out


def compute_lead_time_performance(df: pd.DataFrame, city: str | None = None) -> pd.DataFrame:
    """Hit rate for each model, binned by lead time at observation time."""
    settled = _settled_subset(df, city)
    if settled.empty or "lead_time_hours" not in settled.columns:
        return pd.DataFrame(columns=["lead_bucket", "n"] + [f"{m}_hit_rate" for m in MODEL_HIT_COLS])

    settled = settled.copy()
    settled["lead_bucket"] = pd.cut(
        settled["lead_time_hours"], bins=LEAD_TIME_BINS, labels=LEAD_TIME_LABELS
    )
    rows = []
    for bucket in LEAD_TIME_LABELS:
        group = settled[settled["lead_bucket"] == bucket]
        if group.empty:
            continue
        row = {"lead_bucket": bucket, "n": len(group)}
        for model, col in MODEL_HIT_COLS.items():
            if col in group.columns:
                vals = group[col].dropna()
                row[f"{model}_hit_rate"] = vals.mean() if len(vals) else None
                row[f"{model}_n"] = len(vals)
            else:
                row[f"{model}_hit_rate"] = None
                row[f"{model}_n"] = 0
        rows.append(row)
    return pd.DataFrame(rows)


def compute_city_performance(df: pd.DataFrame) -> pd.DataFrame:
    """One row per city: bucket hit rate for each model, plus the
    "best model" for that city (highest hit rate, ties broken by
    whichever's listed first in MODEL_HIT_COLS). This directly answers
    "which city does each model track best" at the bucket level.
    """
    settled = _settled_subset(df, city=None)
    if settled.empty:
        return pd.DataFrame(columns=["city", "n", "best_model"] + [f"{m}_hit_rate" for m in MODEL_HIT_COLS])

    rows = []
    for city, group in settled.groupby("city"):
        row = {"city": city, "n": len(group)}
        rates = {}
        for model, col in MODEL_HIT_COLS.items():
            if col in group.columns:
                vals = group[col].dropna()
                rate = vals.mean() if len(vals) else None
                row[f"{model}_hit_rate"] = rate
                row[f"{model}_n"] = len(vals)
                if rate is not None:
                    rates[model] = rate
            else:
                row[f"{model}_hit_rate"] = None
                row[f"{model}_n"] = 0
        row["best_model"] = max(rates, key=rates.get) if rates else None
        rows.append(row)
    return pd.DataFrame(rows).sort_values("n", ascending=False)


def compute_city_temp_accuracy(df: pd.DataFrame) -> pd.DataFrame:
    """One row per city: mean absolute error (°C) between each model's
    forecast and the actual settled temperature. This is a finer-grained
    answer to "how closely does this city follow the models" than the
    bucket hit rate above -- a city can just barely miss its bucket
    every time (small MAE, 0% hit rate) or hit the bucket by chance with
    a forecast that was actually far off (larger MAE, high hit rate).
    Look at both together, not just one.

    Only uses rows where actual_max_c_used is known -- does not require
    a resolved bucket match, so this can include rows evaluate_row()
    couldn't map to a bucket.
    """
    if "actual_max_c_used" not in df.columns:
        return pd.DataFrame(columns=["city", "n", "ECMWF_mae_c", "GFS_mae_c"])

    known = df[df["actual_max_c_used"].notna()].copy()
    if known.empty:
        return pd.DataFrame(columns=["city", "n", "ECMWF_mae_c", "GFS_mae_c"])

    rows = []
    for city, group in known.groupby("city"):
        row = {"city": city, "n": len(group)}
        for model, col in {"ECMWF": "ecmwf_max_c", "GFS": "gfs_max_c"}.items():
            if col in group.columns:
                diffs = (group[col] - group["actual_max_c_used"]).abs().dropna()
                row[f"{model}_mae_c"] = diffs.mean() if len(diffs) else None
                row[f"{model}_mae_n"] = len(diffs)
            else:
                row[f"{model}_mae_c"] = None
                row[f"{model}_mae_n"] = 0
        rows.append(row)
    return pd.DataFrame(rows).sort_values("n", ascending=False)



def compute_overall_summary(df: pd.DataFrame, city: str | None = None) -> dict:
    settled = _settled_subset(df, city)
    summary = {"scope": city or "All Cities", "settled_n": len(settled)}
    for model, col in MODEL_HIT_COLS.items():
        if col in settled.columns:
            vals = settled[col].dropna()
            summary[f"{model}_hit_rate"] = vals.mean() if len(vals) else None
            summary[f"{model}_n"] = len(vals)
        else:
            summary[f"{model}_hit_rate"] = None
            summary[f"{model}_n"] = 0
    return summary


def generate_insights(
    summary: dict, daily: pd.DataFrame, lead: pd.DataFrame,
    city_perf: pd.DataFrame | None = None, city_mae: pd.DataFrame | None = None,
) -> list[str]:
    """Small set of auto-generated, plainly-worded observations. These
    are DESCRIPTIVE statements about what's in the data, not trading
    recommendations -- deliberately no "buy/sell"/"edge" language.
    """
    insights = []
    n = summary.get("settled_n", 0)

    if n < 20:
        insights.append(
            f"Only {n} settled observations in this report -- treat every "
            "number below as preliminary. Model comparisons typically need "
            "dozens of resolved outcomes per slice before differences are "
            "distinguishable from noise."
        )

    rates = {m: summary.get(f"{m}_hit_rate") for m in MODEL_HIT_COLS}
    valid_rates = {m: r for m, r in rates.items() if r is not None}
    if len(valid_rates) >= 2:
        best = max(valid_rates, key=valid_rates.get)
        worst = min(valid_rates, key=valid_rates.get)
        if best != worst:
            insights.append(
                f"{best} had the highest bucket hit rate in this sample "
                f"({valid_rates[best]*100:.0f}%), vs {worst} at "
                f"{valid_rates[worst]*100:.0f}%."
            )

    if not lead.empty and "ECMWF_hit_rate" in lead.columns:
        valid = lead.dropna(subset=["ECMWF_hit_rate"])
        if len(valid) >= 2:
            first, last = valid.iloc[0], valid.iloc[-1]
            direction = "higher" if last["ECMWF_hit_rate"] > first["ECMWF_hit_rate"] else "lower"
            insights.append(
                f"ECMWF's hit rate at {last['lead_bucket']} lead time "
                f"({last['ECMWF_hit_rate']*100:.0f}%) was {direction} than "
                f"at {first['lead_bucket']} ({first['ECMWF_hit_rate']*100:.0f}%) "
                "in this sample."
            )

    if summary.get("Market Favorite_hit_rate") is not None:
        insights.append(
            f"The market's own favorite (highest-priced) bucket won "
            f"{summary['Market Favorite_hit_rate']*100:.0f}% of the time "
            f"in this sample -- this is a baseline for how often 'the "
            "crowd' is simply right, independent of either model."
        )

    # --- Per-city findings (only meaningful in the all-cities report) ---
    if city_perf is not None and not city_perf.empty:
        MIN_CITY_N = 10  # don't call out a city's "best model" on a handful of rows
        reliable = city_perf[city_perf["n"] >= MIN_CITY_N].dropna(subset=["best_model"])
        if not reliable.empty:
            counts = reliable["best_model"].value_counts()
            lead_model = counts.idxmax()
            insights.append(
                f"{lead_model} was the best-performing model (by bucket hit "
                f"rate) in {counts[lead_model]} of {len(reliable)} cities "
                f"with at least {MIN_CITY_N} settled observations."
            )
            # Highlight the single strongest and weakest ECMWF city as a
            # concrete example, since that's usually the model people
            # ask about first.
            ecmwf_rates = reliable.dropna(subset=["ECMWF_hit_rate"])
            if len(ecmwf_rates) >= 2:
                best_row = ecmwf_rates.loc[ecmwf_rates["ECMWF_hit_rate"].idxmax()]
                worst_row = ecmwf_rates.loc[ecmwf_rates["ECMWF_hit_rate"].idxmin()]
                if best_row["city"] != worst_row["city"]:
                    insights.append(
                        f"ECMWF's bucket hit rate ranged from "
                        f"{worst_row['ECMWF_hit_rate']*100:.0f}% in {worst_row['city']} "
                        f"to {best_row['ECMWF_hit_rate']*100:.0f}% in {best_row['city']} "
                        f"-- accuracy is not uniform across cities."
                    )
        elif len(city_perf) > 0:
            insights.append(
                f"No single city has {MIN_CITY_N}+ settled observations yet, "
                "so per-city model comparisons aren't reliable yet -- check "
                "back once more data accumulates."
            )

    if city_mae is not None and not city_mae.empty:
        reliable_mae = city_mae[city_mae["n"] >= 10].dropna(subset=["ECMWF_mae_c", "GFS_mae_c"])
        if not reliable_mae.empty:
            tightest = reliable_mae.loc[
                reliable_mae[["ECMWF_mae_c", "GFS_mae_c"]].min(axis=1).idxmin()
            ]
            insights.append(
                f"{tightest['city']} had the tightest forecast-to-actual "
                f"temperature gap in this sample (ECMWF off by "
                f"{tightest['ECMWF_mae_c']:.1f}°C, GFS by "
                f"{tightest['GFS_mae_c']:.1f}°C on average) -- note this is "
                "a temperature-error view, which can disagree with the "
                "bucket-hit-rate view above for cities near a bucket edge."
            )

    if not insights:
        insights.append("Not enough settled data yet to generate insights.")

    return insights


# ---------------------------------------------------------------------------
# Chart helpers (shared between PNG and DOCX paths)
# ---------------------------------------------------------------------------
_COLORS = {"ECMWF": "#4C72B0", "GFS": "#DD8452", "Market Favorite": "#55A868"}


def _plot_hit_rate_by(ax, table: pd.DataFrame, x_col: str, title: str):
    if table.empty:
        ax.text(0.5, 0.5, "Not enough data", ha="center", va="center", transform=ax.transAxes)
        ax.set_title(title)
        return
    x = range(len(table))
    width = 0.25
    for i, model in enumerate(MODEL_HIT_COLS):
        rate_col = f"{model}_hit_rate"
        if rate_col not in table.columns:
            continue
        offsets = [xi + (i - 1) * width for xi in x]
        ax.bar(offsets, table[rate_col].fillna(0), width=width, label=model, color=_COLORS[model])
    ax.set_xticks(list(x))
    ax.set_xticklabels(table[x_col].astype(str), rotation=30, ha="right")
    ax.set_ylim(0, 1)
    ax.set_ylabel("Hit Rate")
    ax.set_title(title)
    ax.legend(fontsize=8)


def _plot_city_comparison(ax, city_table: pd.DataFrame, title: str, max_cities: int = 14):
    if city_table.empty:
        ax.text(0.5, 0.5, "Not enough data", ha="center", va="center", transform=ax.transAxes)
        ax.set_title(title)
        return
    table = city_table.head(max_cities).sort_values("city")
    x = range(len(table))
    width = 0.25
    for i, model in enumerate(MODEL_HIT_COLS):
        rate_col = f"{model}_hit_rate"
        if rate_col not in table.columns:
            continue
        offsets = [xi + (i - 1) * width for xi in x]
        ax.bar(offsets, table[rate_col].fillna(0), width=width, label=model, color=_COLORS[model])
    ax.set_xticks(list(x))
    ax.set_xticklabels(table["city"].astype(str), rotation=45, ha="right")
    ax.set_ylim(0, 1)
    ax.set_ylabel("Hit Rate")
    ax.set_title(title)
    ax.legend(fontsize=8)


def _render_charts_figure(daily: pd.DataFrame, lead: pd.DataFrame, scope_label: str) -> plt.Figure:
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.2))
    _plot_hit_rate_by(axes[0], daily, "target_date", f"Hit Rate by Day - {scope_label}")
    _plot_hit_rate_by(axes[1], lead, "lead_bucket", f"Hit Rate by Lead Time - {scope_label}")
    fig.tight_layout()
    return fig


# ---------------------------------------------------------------------------
# PNG report
# ---------------------------------------------------------------------------
def build_png_report(df: pd.DataFrame, city: str | None = None) -> bytes:
    scope_label = city or "All Cities"
    summary = compute_overall_summary(df, city)
    daily = compute_daily_performance(df, city)
    lead = compute_lead_time_performance(df, city)
    city_perf = compute_city_performance(df) if city is None else None
    city_mae = compute_city_temp_accuracy(df) if city is None else None
    insights = generate_insights(summary, daily, lead, city_perf, city_mae)

    show_city_panel = city is None and city_perf is not None and not city_perf.empty

    fig = plt.figure(figsize=(11, 11.5 if show_city_panel else 7.5))
    n_rows = 4 if show_city_panel else 3
    height_ratios = [1.0, 2.4, 2.6, 1.6] if show_city_panel else [0.7, 2.4, 1.1]
    gs = fig.add_gridspec(n_rows, 2, height_ratios=height_ratios, hspace=0.7)

    # Header
    ax_header = fig.add_subplot(gs[0, :])
    ax_header.axis("off")
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    header_text = (
        f"Weather Market Model Performance Report -- {scope_label}\n"
        f"Generated {generated}  |  Settled observations: {summary['settled_n']}"
    )
    ax_header.text(0, 0.95, header_text, fontsize=13, fontweight="bold", va="top")
    hit_rate_line = "   ".join(
        f"{m}: {summary[f'{m}_hit_rate']*100:.0f}% (n={summary[f'{m}_n']})"
        if summary.get(f"{m}_hit_rate") is not None else f"{m}: n/a"
        for m in MODEL_HIT_COLS
    )
    ax_header.text(0, 0.15, hit_rate_line, fontsize=10, va="top")

    ax1 = fig.add_subplot(gs[1, 0])
    _plot_hit_rate_by(ax1, daily, "target_date", "Hit Rate by Day")
    ax2 = fig.add_subplot(gs[1, 1])
    _plot_hit_rate_by(ax2, lead, "lead_bucket", "Hit Rate by Lead Time")

    insights_row = 2
    if show_city_panel:
        ax_city = fig.add_subplot(gs[2, :])
        _plot_city_comparison(ax_city, city_perf, "Hit Rate by City (which city each model tracks best)")
        insights_row = 3

    ax_insights = fig.add_subplot(gs[insights_row, :])
    ax_insights.axis("off")
    insight_text = "Insights:\n" + "\n".join(f"- {i}" for i in insights)
    ax_insights.text(0, 1.0, insight_text, fontsize=9, va="top", wrap=True)

    caveat = (
        "Settlement uses official station data where available (NOAA/HKO), "
        "an Open-Meteo reanalysis proxy otherwise -- see SCHEMA.md. "
        "This report is descriptive, not a trading signal."
    )
    fig.text(0.02, 0.005, caveat, fontsize=7, color="gray")

    fig.tight_layout(rect=[0, 0.02, 1, 0.98])
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX report
# ---------------------------------------------------------------------------
def build_docx_report(df: pd.DataFrame, city: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    scope_label = city or "All Cities"
    summary = compute_overall_summary(df, city)
    daily = compute_daily_performance(df, city)
    lead = compute_lead_time_performance(df, city)
    city_perf = compute_city_performance(df) if city is None else None
    city_mae = compute_city_temp_accuracy(df) if city is None else None
    insights = generate_insights(summary, daily, lead, city_perf, city_mae)

    doc = Document()

    title = doc.add_heading(f"Weather Market Model Performance Report", level=1)
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"Scope: {scope_label}").bold = True
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_paragraph(f"Generated: {generated}")
    doc.add_paragraph(f"Settled observations included: {summary['settled_n']}")

    doc.add_paragraph(
        "Settlement uses official station data where available (NOAA/NWS for "
        "New York, Chicago, Miami; Hong Kong Observatory for Hong Kong), and "
        "an Open-Meteo reanalysis proxy for all other cities. See SCHEMA.md "
        "for details. This report is descriptive only -- it is not a trading "
        "recommendation."
    )

    doc.add_heading("Overall Hit Rates", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Model", "Hit Rate", "Sample Size (n)"
    for model in MODEL_HIT_COLS:
        row = table.add_row().cells
        rate = summary.get(f"{model}_hit_rate")
        row[0].text = model
        row[1].text = f"{rate*100:.1f}%" if rate is not None else "n/a"
        row[2].text = str(summary.get(f"{model}_n", 0))

    doc.add_heading("Charts", level=2)
    fig = _render_charts_figure(daily, lead, scope_label)
    img_buf = io.BytesIO()
    fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(6.3))

    doc.add_heading("Hit Rate by Day", level=2)
    if daily.empty:
        doc.add_paragraph("No settled data available for this scope yet.")
    else:
        t = doc.add_table(rows=1, cols=1 + len(MODEL_HIT_COLS) + 1)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text = "Target Date"
        for i, model in enumerate(MODEL_HIT_COLS, start=1):
            hdr[i].text = f"{model} Hit Rate"
        hdr[-1].text = "n"
        for _, r in daily.iterrows():
            cells = t.add_row().cells
            cells[0].text = str(r["target_date"])
            for i, model in enumerate(MODEL_HIT_COLS, start=1):
                rate = r.get(f"{model}_hit_rate")
                cells[i].text = f"{rate*100:.0f}%" if pd.notnull(rate) else "n/a"
            cells[-1].text = str(int(r["n"]))

    doc.add_heading("Hit Rate by Lead Time", level=2)
    if lead.empty:
        doc.add_paragraph("No settled data available for this scope yet.")
    else:
        t2 = doc.add_table(rows=1, cols=1 + len(MODEL_HIT_COLS) + 1)
        t2.style = "Light Grid Accent 1"
        hdr = t2.rows[0].cells
        hdr[0].text = "Lead Time"
        for i, model in enumerate(MODEL_HIT_COLS, start=1):
            hdr[i].text = f"{model} Hit Rate"
        hdr[-1].text = "n"
        for _, r in lead.iterrows():
            cells = t2.add_row().cells
            cells[0].text = str(r["lead_bucket"])
            for i, model in enumerate(MODEL_HIT_COLS, start=1):
                rate = r.get(f"{model}_hit_rate")
                cells[i].text = f"{rate*100:.0f}%" if pd.notnull(rate) else "n/a"
            cells[-1].text = str(int(r["n"]))

    # --- Per-city breakdown: the key "which city follows which model"
    # view -- only meaningful in the all-cities report; a single-city
    # report is already scoped to one city.
    if city is None and city_perf is not None and not city_perf.empty:
        doc.add_heading("Performance by City", level=2)
        doc.add_paragraph(
            "Bucket hit rate per city, per model. 'Best Model' is whichever "
            "had the highest hit rate for that city (blank if fewer than "
            "10 settled observations -- too little data to call it)."
        )

        fig_city, ax_city = plt.subplots(figsize=(9, 4))
        _plot_city_comparison(ax_city, city_perf, "Hit Rate by City")
        city_img_buf = io.BytesIO()
        fig_city.savefig(city_img_buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig_city)
        city_img_buf.seek(0)
        doc.add_picture(city_img_buf, width=Inches(6.3))

        t3 = doc.add_table(rows=1, cols=2 + len(MODEL_HIT_COLS) + 1)
        t3.style = "Light Grid Accent 1"
        hdr = t3.rows[0].cells
        hdr[0].text = "City"
        for i, model in enumerate(MODEL_HIT_COLS, start=1):
            hdr[i].text = f"{model} Hit Rate"
        hdr[-2].text = "n"
        hdr[-1].text = "Best Model"
        for _, r in city_perf.iterrows():
            cells = t3.add_row().cells
            cells[0].text = str(r["city"])
            for i, model in enumerate(MODEL_HIT_COLS, start=1):
                rate = r.get(f"{model}_hit_rate")
                cells[i].text = f"{rate*100:.0f}%" if pd.notnull(rate) else "n/a"
            cells[-2].text = str(int(r["n"]))
            cells[-1].text = str(r.get("best_model") or "-") if r["n"] >= 10 else "-"

        if city_mae is not None and not city_mae.empty:
            doc.add_heading("Forecast Temperature Accuracy by City (\u00b0C)", level=3)
            doc.add_paragraph(
                "Mean absolute error between each model's forecast and the "
                "actual settled temperature -- a finer-grained view than "
                "bucket hit rate above. A city can narrowly miss its "
                "bucket every time (small error, low hit rate) or hit by "
                "chance with a forecast that was actually far off."
            )
            t4 = doc.add_table(rows=1, cols=3)
            t4.style = "Light Grid Accent 1"
            hdr = t4.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "City", "ECMWF MAE (\u00b0C)", "GFS MAE (\u00b0C)"
            for _, r in city_mae.iterrows():
                cells = t4.add_row().cells
                cells[0].text = str(r["city"])
                ecmwf_mae = r.get("ECMWF_mae_c")
                gfs_mae = r.get("GFS_mae_c")
                cells[1].text = f"{ecmwf_mae:.1f}" if pd.notnull(ecmwf_mae) else "n/a"
                cells[2].text = f"{gfs_mae:.1f}" if pd.notnull(gfs_mae) else "n/a"

    doc.add_heading("Insights", level=2)
    for insight in insights:
        doc.add_paragraph(insight, style="List Bullet")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
